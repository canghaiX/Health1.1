import docx
import re
from pathlib import Path

def docx_to_markdown(docx_path, md_path=None):
    docx_path = Path(docx_path)
    # 若未指定输出路径，则在同目录下生成同名但扩展名为 .md 的文件
    if md_path is None:
        md_path = docx_path.with_name(f"{docx_path.stem}.md")  # 关键修改点
    # md_path.parent.mkdir(parents=True, exist_ok=True)
    
    doc = docx.Document(docx_path)
    markdown_content = []
    
    for paragraph in doc.paragraphs:
        # 处理标题
        if paragraph.style.name.startswith('Heading'):
            level = int(paragraph.style.name.split(' ')[-1])
            heading = '#' * level + ' ' + paragraph.text
            markdown_content.append(heading)
        
        # 处理列表
        elif paragraph.style.name.startswith('List'):
            # 识别列表类型（有序/无序列表）
            if paragraph.style.name.startswith('List Bullet'):
                bullet = '- '
            elif paragraph.style.name.startswith('List Number'):
                # 简化处理有序列表编号
                bullet = '1. '
            else:
                bullet = '- '  # 默认使用无序列表
            
            list_item = bullet + paragraph.text
            markdown_content.append(list_item)
        
        # 处理普通段落
        else:
            # 处理段落中的内联样式（粗体、斜体）
            text = ''
            for run in paragraph.runs:
                run_text = run.text
                if run.bold:
                    run_text = f'**{run_text}**'
                if run.italic:
                    run_text = f'*{run_text}*'
                text += run_text
            
            # 处理超链接
            text = re.sub(r'<a href="(.*?)">(.*?)</a>', r'[\2](\1)', text)
            
            markdown_content.append(text)
    
    # 处理表格
    for table in doc.tables:
        markdown_table = []
        # 生成Markdown表格的头部和分隔线
        header = []
        for cell in table.rows[0].cells:
            header.append(cell.text.strip())
        markdown_table.append('| ' + ' | '.join(header) + ' |')
        markdown_table.append('| ' + ' | '.join(['---'] * len(header)) + ' |')
        
        # 生成表格内容
        for row in table.rows[1:]:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            markdown_table.append('| ' + ' | '.join(row_data) + ' |')
        
        markdown_content.extend(markdown_table)
        markdown_content.append('')  # 添加空行分隔表格
    
    # 将结果写入Markdown文件
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(markdown_content))
    return md_path,markdown_content

# 使用示例
if __name__ == "__main__":
    docx_file = "example.docx"  # 替换为你的Word文件路径
    md_file = docx_to_markdown(docx_file)
    print(f"已将Word文件转换为Markdown文件：{md_file}")